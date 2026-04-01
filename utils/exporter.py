"""Export functionality for Word (.docx) and PDF formats."""

import io
from datetime import datetime
from typing import List, Dict

def export_to_word(qa_items: List[Dict], institution: str, question_type: str) -> bytes:
    """Export questions and answers to Word document."""
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    
    doc = Document()
    
    # Title
    title = doc.add_heading(f'{institution} - Question Bank', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Metadata
    doc.add_paragraph(f'Generated: {datetime.now().strftime("%B %d, %Y at %I:%M %p")}')
    doc.add_paragraph(f'Question Type: {question_type}')
    doc.add_paragraph('_' * 60)
    
    # Questions and Answers
    for idx, item in enumerate(qa_items, 1):
        # Question
        q_para = doc.add_paragraph()
        q_run = q_para.add_run(f'Q{idx}. {item.get("question", "")}')
        q_run.bold = True
        q_run.font.size = Pt(11)
        q_run.font.color.rgb = RGBColor(26, 58, 92)  # Academic blue
        
        # Answer
        a_para = doc.add_paragraph()
        a_run = a_para.add_run(f'Ans: {item.get("answer", "")}')
        a_run.font.size = Pt(10)
        a_run.font.color.rgb = RGBColor(44, 62, 80)
        
        # Topic and concepts if available
        if item.get("topic"):
            doc.add_paragraph(f'Topic: {item["topic"]}', style='Intense Quote')
        
        if item.get("key_concepts"):
            concepts = ', '.join(item["key_concepts"])
            doc.add_paragraph(f'Key Concepts: {concepts}', style='Quote')
        
        doc.add_paragraph()  # Spacer
    
    # Footer
    section = doc.sections[0]
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.text = f"AI-Powered Question Bank | {institution} | Confidential"
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Save to bytes
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()


def export_to_pdf(qa_items: List[Dict], institution: str, question_type: str) -> bytes:
    """Export questions and answers to PDF document."""
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        rightMargin=0.75*inch,
        leftMargin=0.75*inch,
        topMargin=0.75*inch,
        bottomMargin=0.75*inch
    )
    
    styles = getSampleStyleSheet()
    
    # Custom styles
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=18,
        textColor=colors.HexColor('#1a3a5c'),
        alignment=1,  # Center
        spaceAfter=20
    )
    
    question_style = ParagraphStyle(
        'Question',
        parent=styles['Normal'],
        fontSize=11,
        textColor=colors.HexColor('#1a3a5c'),
        fontName='Helvetica-Bold',
        spaceBefore=12,
        spaceAfter=6
    )
    
    answer_style = ParagraphStyle(
        'Answer',
        parent=styles['Normal'],
        fontSize=10,
        textColor=colors.HexColor('#2c3e50'),
        leftIndent=20,
        spaceAfter=15
    )
    
    # Build content
    story = []
    
    # Title
    story.append(Paragraph(f'{institution}<br/><font size="14">Question Bank</font>', title_style))
    story.append(Spacer(1, 0.2*inch))
    
    # Metadata table
    meta_data = [
        ['Generated:', datetime.now().strftime("%B %d, %Y")],
        ['Question Type:', question_type],
        ['Academic Level:', qa_items[0].get('academic_level', 'N/A') if qa_items else 'N/A']
    ]
    meta_table = Table(meta_data, colWidths=[2*inch, 3*inch])
    meta_table.setStyle(TableStyle([
        ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, -1), 9),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
        ('BACKGROUND', (0, 0), (0, -1), colors.HexColor('#f8f9fa')),
        ('PADDING', (0, 0), (-1, -1), 6),
    ]))
    story.append(meta_table)
    story.append(Spacer(1, 0.3*inch))
    
    # Questions and Answers
    for idx, item in enumerate(qa_items, 1):
        # Question
        story.append(Paragraph(f'<b>Q{idx}.</b> {item.get("question", "")}', question_style))
        
        # Answer
        answer_text = item.get("answer", "No answer generated")
        story.append(Paragraph(f'<i>Answer:</i> {answer_text}', answer_style))
        
        # Additional info
        if item.get("topic") or item.get("key_concepts"):
            extra_info = []
            if item.get("topic"):
                extra_info.append(f"Topic: {item['topic']}")
            if item.get("key_concepts"):
                extra_info.append(f"Concepts: {', '.join(item['key_concepts'][:3])}")
            story.append(Paragraph(f'<font color="grey"><i>{" | ".join(extra_info)}</i></font>', 
                                 ParagraphStyle('Meta', parent=styles['Normal'], fontSize=8)))
        
        story.append(Spacer(1, 0.15*inch))
    
    # Build PDF
    doc.build(story)
    buffer.seek(0)
    return buffer.read()
